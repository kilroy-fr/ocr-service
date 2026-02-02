"""
File Routes - Upload, Download, Preview
"""
import os
import time
from flask import Blueprint, request, jsonify, send_from_directory
from werkzeug.utils import secure_filename
from urllib.parse import unquote
from html import escape

from services.logger import log
from services.session_manager import ensure_staging
from services.file_utils import fs
from config import INPUT_ROOT

file_bp = Blueprint('file', __name__)


@file_bp.route('/upload', methods=['POST'])
def upload_files():
    """
    Einzel-Upload: Dateien NUR ins Staging kopieren (OHNE OCR/Analyse).
    OCR/Analyse erfolgt erst beim Klick auf "Analysieren!" (copy_and_analyze).
    """
    session_id = ensure_staging()

    files = request.files.getlist('files')
    if not files or all(f.filename == '' for f in files):
        log("Keine Dateien ausgewählt")
        return jsonify(success=False, message="Keine Dateien ausgewählt"), 400

    # Dateien NUR ins Staging kopieren (kein OCR)
    uploaded_files = []
    for file in files:
        if file.filename:
            from PIL import Image
            filename = file.filename

            # Prüfe ob Datei ein TIFF ist (durch Magic Bytes, nicht durch Endung)
            # Lese die ersten 4 Bytes um TIFF zu erkennen
            file.seek(0)
            magic_bytes = file.read(4)
            file.seek(0)  # Zurück zum Anfang

            is_tiff = magic_bytes[:4] in (b'II*\x00', b'MM\x00*')  # TIFF Little/Big Endian

            # ✅ Dateien ohne gültige Dateiendung behandeln
            has_valid_extension = filename.lower().endswith(('.pdf', '.jpg', '.jpeg', '.png', '.tif', '.tiff', '.txt', '.docx'))

            if not has_valid_extension:
                if is_tiff:
                    # TIFF ohne Endung → .tif ergänzen
                    log(f"📝 TIFF-Datei ohne Endung erkannt: {filename}")
                    filename = filename + '.tif'
                    log(f"✅ TIFF-Endung ergänzt: {filename}")
                else:
                    # Andere Datei ohne Endung → .jpg ergänzen
                    log(f"📝 Datei ohne gültige Endung erkannt: {filename}")
                    filename = filename + '.jpg'
                    log(f"✅ JPG-Endung ergänzt: {filename}")

            # Sichere den Dateinamen
            safe_name = secure_filename(filename)

            # ✅ TIF/TIFF direkt zu JPG konvertieren (durch Magic Bytes, nicht durch Endung)
            if is_tiff or safe_name.lower().endswith(('.tif', '.tiff')):
                try:
                    log(f"🔄 TIF/TIFF wird zu JPG konvertiert: {safe_name}")

                    # TIF direkt aus Upload-Stream öffnen
                    img = Image.open(file.stream)

                    # Erste Seite bei Multi-Page TIFF
                    if hasattr(img, 'n_frames') and img.n_frames > 1:
                        img.seek(0)

                    # RGB konvertieren für JPG
                    if img.mode not in ('RGB', 'L'):
                        img = img.convert('RGB')

                    # Neuer JPG-Dateiname
                    base_name = os.path.splitext(safe_name)[0]
                    jpg_name = base_name + '.jpg'
                    jpg_path = os.path.join(fs.work_dir, jpg_name)

                    # Als JPG speichern
                    img.save(jpg_path, 'JPEG', quality=95)
                    img.close()

                    log(f"✅ TIF zu JPG konvertiert (TIF nicht gespeichert): {jpg_name}")
                    uploaded_files.append(jpg_name)
                except Exception as e:
                    log(f"❌ Fehler bei TIF-Konvertierung: {e}", level="error")
                    # Bei Fehler: TIF-Original speichern
                    staging_path = os.path.join(fs.work_dir, safe_name)
                    file.seek(0)  # Stream zurückspulen
                    file.save(staging_path)
                    uploaded_files.append(safe_name)
            else:
                # Normale Dateien (PDF, JPG, PNG) direkt speichern
                staging_path = os.path.join(fs.work_dir, safe_name)
                file.save(staging_path)
                uploaded_files.append(safe_name)

            log(f"📤 Datei ins Staging kopiert: {uploaded_files[-1]}")

    if not uploaded_files:
        return jsonify(success=False, message="Keine Dateien hochgeladen"), 400

    log(f"✅ {len(uploaded_files)} Datei(en) ins Staging kopiert (ohne OCR)")
    return jsonify(success=True, files=uploaded_files, count=len(uploaded_files))


@file_bp.route('/upload_folder', methods=['POST'])
def upload_folder():
    """
    Batch-Upload: Dateien NUR ins Staging kopieren (OHNE OCR/Analyse).
    OCR/Analyse erfolgt erst beim Klick auf "Analysieren!" (copy_and_analyze).
    """
    session_id = ensure_staging()

    uploaded_files = request.files.getlist("files")
    if not uploaded_files or all(f.filename == '' for f in uploaded_files):
        log("Keine Dateien empfangen")
        return jsonify(success=False, message="Keine Dateien empfangen"), 400

    # Alle Dateien sammeln (flach, ohne Ordnerstruktur)
    collected_files = []

    for file in uploaded_files:
        if not file.filename:
            continue

        from PIL import Image

        # Nur Dateiname, keine Pfadstruktur
        filename = os.path.basename(file.filename)

        # Prüfe ob Datei ein TIFF ist (durch Magic Bytes, nicht durch Endung)
        file.seek(0)
        magic_bytes = file.read(4)
        file.seek(0)

        is_tiff = magic_bytes[:4] in (b'II*\x00', b'MM\x00*')

        # ✅ Dateien ohne gültige Dateiendung behandeln
        has_valid_extension = filename.lower().endswith(('.pdf', '.jpg', '.jpeg', '.png', '.tif', '.tiff', '.txt'))

        if not has_valid_extension:
            if is_tiff:
                # TIFF ohne Endung → .tif ergänzen
                log(f"📝 TIFF-Datei ohne Endung erkannt: {filename}")
                filename = filename + '.tif'
                log(f"✅ TIFF-Endung ergänzt: {filename}")
            else:
                # Andere Datei ohne Endung → .jpg ergänzen
                log(f"📝 Datei ohne gültige Endung erkannt: {filename}")
                filename = filename + '.jpg'
                log(f"✅ JPG-Endung ergänzt: {filename}")

        # Sichere den Dateinamen
        safe_name = secure_filename(filename)

        # Nur unterstützte Formate
        if not safe_name.lower().endswith(('.pdf', '.jpg', '.jpeg', '.png', '.tif', '.tiff', '.txt', '.docx')):
            log(f"⏭️ Überspringe Datei (nicht unterstützt): {safe_name}")
            continue

        # ✅ TIF/TIFF direkt zu JPG konvertieren (durch Magic Bytes, nicht durch Endung)
        if is_tiff or safe_name.lower().endswith(('.tif', '.tiff')):
            try:
                log(f"🔄 TIF/TIFF wird zu JPG konvertiert: {safe_name}")

                # TIF direkt aus Upload-Stream öffnen
                img = Image.open(file.stream)

                # Erste Seite bei Multi-Page TIFF
                if hasattr(img, 'n_frames') and img.n_frames > 1:
                    img.seek(0)

                # RGB konvertieren für JPG
                if img.mode not in ('RGB', 'L'):
                    img = img.convert('RGB')

                # Neuer JPG-Dateiname
                base_name = os.path.splitext(safe_name)[0]
                jpg_name = base_name + '.jpg'
                jpg_path = os.path.join(fs.work_dir, jpg_name)

                # Bei Namenskonflikten: Timestamp hinzufügen
                if os.path.exists(jpg_path):
                    jpg_name = f"{base_name}_{int(time.time())}.jpg"
                    jpg_path = os.path.join(fs.work_dir, jpg_name)

                # Als JPG speichern
                img.save(jpg_path, 'JPEG', quality=95)
                img.close()

                log(f"✅ TIF zu JPG konvertiert (TIF nicht gespeichert): {jpg_name}")
                collected_files.append(jpg_name)
            except Exception as e:
                log(f"❌ Fehler bei TIF-Konvertierung: {e}", level="error")
                # Bei Fehler: TIF-Original speichern
                staging_path = os.path.join(fs.work_dir, safe_name)
                if os.path.exists(staging_path):
                    base, ext = os.path.splitext(safe_name)
                    safe_name = f"{base}_{int(time.time())}{ext}"
                    staging_path = os.path.join(fs.work_dir, safe_name)
                file.seek(0)  # Stream zurückspulen
                file.save(staging_path)
                collected_files.append(safe_name)
        else:
            # Normale Dateien (PDF, JPG, PNG) direkt speichern
            staging_path = os.path.join(fs.work_dir, safe_name)

            # Bei Namenskonflikten: Präfix mit Timestamp
            if os.path.exists(staging_path):
                base, ext = os.path.splitext(safe_name)
                safe_name = f"{base}_{int(time.time())}{ext}"
                staging_path = os.path.join(fs.work_dir, safe_name)

            file.save(staging_path)
            collected_files.append(safe_name)

        log(f"📤 Datei ins Staging kopiert: {collected_files[-1]}")

    if not collected_files:
        return jsonify(success=False, message="Keine verarbeitbaren Dateien gefunden"), 400

    log(f"📁 Batch-Upload: {len(collected_files)} Datei(en) ins Staging kopiert (ohne OCR)")
    return jsonify(success=True, files=collected_files, count=len(collected_files))


@file_bp.route('/preview/<path:filename>')
def preview_file(filename):
    """Vorschau - TIF-Dateien wurden bereits beim Upload zu JPG konvertiert."""
    filename = unquote(filename)

    log(f"🔍 [DEBUG] Preview-Request für: {filename}")
    log(f"🔍 [DEBUG] Session-ID: {fs.session_id if hasattr(fs, 'session_id') else 'keine'}")

    # Sicherer Zugriff auf work_dir
    work_dir_str = None
    if fs.session_id:
        try:
            work_dir_str = str(fs.work_dir)
            log(f"🔍 [DEBUG] work_dir: {work_dir_str}")
        except Exception as e:
            log(f"⚠️ Fehler beim Zugriff auf work_dir: {e}", level="warning")

    # Pfad zur Datei finden
    file_path = None
    if work_dir_str:
        staged = os.path.join(work_dir_str, filename)
        log(f"🔍 [DEBUG] Suche in Staging: {staged}")
        log(f"🔍 [DEBUG] Datei existiert: {os.path.exists(staged)}")

        if os.path.exists(staged):
            log(f"✅ Vorschau-Datei gefunden in Staging: {filename}")
            file_path = staged

    # Fallback: Original im INPUT_ROOT
    if not file_path:
        original = os.path.join(INPUT_ROOT, filename)
        log(f"🔍 [DEBUG] Suche in INPUT_ROOT: {original}")
        log(f"🔍 [DEBUG] Datei existiert: {os.path.exists(original)}")

        if os.path.exists(original):
            log(f"✅ Vorschau-Datei gefunden in INPUT_ROOT: {filename}")
            file_path = original

    if not file_path:
        log(f"⚠️ Vorschau-Datei nicht gefunden: {filename}", level="warning")
        return f"<div style='padding:20px;text-align:center;color:#999;'><h3>❌ Datei nicht gefunden</h3><p>{filename}</p></div>", 404

    # ✅ DOCX-Dateien als HTML mit formatiertem Text zurückgeben
    if filename.lower().endswith('.docx'):
        try:
            from docx import Document

            doc = Document(file_path)

            # Text aus allen Paragraphen sammeln
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(escape(para.text))

            # Text aus Tabellen extrahieren
            for table in doc.tables:
                table_html = ["<table style='border-collapse: collapse; margin: 10px 0;'>"]
                for row in table.rows:
                    table_html.append("<tr>")
                    for cell in row.cells:
                        cell_text = escape(cell.text)
                        table_html.append(f"<td style='border: 1px solid #555; padding: 5px;'>{cell_text}</td>")
                    table_html.append("</tr>")
                table_html.append("</table>")
                paragraphs.append("".join(table_html))

            content_html = "<br><br>".join(paragraphs)

            html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <title>{escape(filename)}</title>
                <style>
                    body {{
                        margin: 0;
                        padding: 20px;
                        font-family: Arial, sans-serif;
                        background: #1e1e1e;
                        color: #e0e0e0;
                        line-height: 1.6;
                    }}
                    table {{
                        width: 100%;
                    }}
                </style>
            </head>
            <body>
                {content_html}
            </body>
            </html>
            """
            log(f"✅ DOCX-Vorschau generiert: {filename}")
            return html, 200, {'Content-Type': 'text/html; charset=utf-8'}
        except Exception as e:
            log(f"❌ Fehler beim Lesen der DOCX-Datei: {e}", level="error")
            return f"<div style='padding:20px;text-align:center;color:#999;'><h3>❌ Fehler beim Lesen</h3><p>{escape(str(e))}</p></div>", 500

    # ✅ TXT-Dateien als HTML mit <pre> zurückgeben
    if filename.lower().endswith('.txt'):
        try:
            # Versuche verschiedene Encodings
            content = None
            encodings = ['utf-8', 'windows-1252', 'latin-1', 'iso-8859-1', 'cp1252']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    log(f"✅ TXT-Datei erfolgreich gelesen mit {encoding}: {filename}")
                    break
                except (UnicodeDecodeError, LookupError):
                    continue

            if content is None:
                # Fallback: Als Binär lesen und ersetzen
                with open(file_path, 'rb') as f:
                    raw_content = f.read()
                content = raw_content.decode('utf-8', errors='replace')
                log(f"⚠️ TXT-Datei mit errors='replace' gelesen: {filename}")

            # HTML-Escape für sichere Darstellung
            escaped_content = escape(content)

            html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <title>{escape(filename)}</title>
                <style>
                    body {{
                        margin: 0;
                        padding: 20px;
                        font-family: 'Courier New', monospace;
                        background: #1e1e1e;
                        color: #e0e0e0;
                    }}
                    pre {{
                        margin: 0;
                        white-space: pre-wrap;
                        word-wrap: break-word;
                        font-size: 14px;
                        line-height: 1.5;
                    }}
                </style>
            </head>
            <body>
                <pre>{escaped_content}</pre>
            </body>
            </html>
            """
            return html, 200, {'Content-Type': 'text/html; charset=utf-8'}
        except Exception as e:
            log(f"❌ Fehler beim Lesen der TXT-Datei: {e}", level="error")
            return f"<div style='padding:20px;text-align:center;color:#999;'><h3>❌ Fehler beim Lesen</h3><p>{escape(str(e))}</p></div>", 500

    # Normale Dateien (Bilder, PDF)
    directory = os.path.dirname(file_path)
    basename = os.path.basename(file_path)
    return send_from_directory(directory, basename)


@file_bp.route('/processed/<path:filename>')
def serve_processed_file(filename):
    """Liefert verarbeitete Dateien aus Staging oder Output."""
    from config import OUTPUT_ROOT
    from flask import make_response

    filename = unquote(filename)

    # Debug-Logging
    log(f"🔍 [DEBUG] /processed angefordert: {filename}")
    log(f"🔍 [DEBUG] Session-ID: {fs.session_id if hasattr(fs, 'session_id') else 'keine'}")

    # Sicherer Zugriff auf work_dir
    work_dir_str = None
    if fs.session_id:
        try:
            work_dir_str = str(fs.work_dir)
        except Exception as e:
            log(f"⚠️ Fehler beim Zugriff auf work_dir: {e}", level="warning")

    if work_dir_str:
        staged = os.path.join(work_dir_str, filename)
        log(f"🔍 [DEBUG] Prüfe Staging: {staged}")
        log(f"🔍 [DEBUG] Staging existiert: {os.path.exists(staged)}")

        if os.path.exists(staged):
            log(f"✅ Datei gefunden in Staging: {filename}")
            # PDF inline anzeigen statt herunterladen
            response = make_response(send_from_directory(work_dir_str, filename))
            if filename.lower().endswith('.pdf'):
                response.headers['Content-Disposition'] = 'inline'
            return response

        # Prüfe auch, ob Datei direkt im work_dir liegt (ohne Unterordner)
        base_name = os.path.basename(filename)
        if base_name != filename:
            staged_base = os.path.join(work_dir_str, base_name)
            log(f"🔍 [DEBUG] Prüfe Staging (nur Dateiname): {staged_base}")
            log(f"🔍 [DEBUG] Staging (Dateiname) existiert: {os.path.exists(staged_base)}")

            if os.path.exists(staged_base):
                log(f"✅ Datei gefunden in Staging (Dateiname): {base_name}")
                # PDF inline anzeigen statt herunterladen
                response = make_response(send_from_directory(work_dir_str, base_name))
                if base_name.lower().endswith('.pdf'):
                    response.headers['Content-Disposition'] = 'inline'
                return response

    output = os.path.join(OUTPUT_ROOT, filename)
    log(f"🔍 [DEBUG] Prüfe Output: {output}")
    log(f"🔍 [DEBUG] Output existiert: {os.path.exists(output)}")

    if os.path.exists(output):
        log(f"✅ Datei gefunden in Output: {filename}")
        # PDF inline anzeigen statt herunterladen
        response = make_response(send_from_directory(OUTPUT_ROOT, filename))
        if filename.lower().endswith('.pdf'):
            response.headers['Content-Disposition'] = 'inline'
        return response

    # Auch im Output nur nach Dateinamen suchen
    base_name = os.path.basename(filename)
    if base_name != filename:
        output_base = os.path.join(OUTPUT_ROOT, base_name)
        log(f"🔍 [DEBUG] Prüfe Output (nur Dateiname): {output_base}")
        log(f"🔍 [DEBUG] Output (Dateiname) existiert: {os.path.exists(output_base)}")

        if os.path.exists(output_base):
            log(f"✅ Datei gefunden in Output (Dateiname): {base_name}")
            # PDF inline anzeigen statt herunterladen
            response = make_response(send_from_directory(OUTPUT_ROOT, base_name))
            if base_name.lower().endswith('.pdf'):
                response.headers['Content-Disposition'] = 'inline'
            return response

    log(f"❌ Verarbeitete Datei nicht gefunden: {filename}", level="error")
    staging_path = os.path.join(work_dir_str, filename) if work_dir_str else 'keine Session'
    log(f"   Gesucht in Staging: {staging_path}", level="error")
    log(f"   Gesucht in Output: {output}", level="error")

    # Gib einen 404 HTML-Response statt JSON zurück
    return f"<div style='padding:20px;text-align:center;color:#999;'><h3>❌ Datei nicht gefunden</h3><p>{filename}</p><p style='font-size:0.8em;margin-top:20px;'>Staging: {staging_path}<br>Output: {output}</p></div>", 404


@file_bp.route("/rotate_file", methods=["POST"])
def rotate_file():
    """
    Rotiert eine Datei (PDF oder Bild) vor der Analyse.
    Unterstützt 90°, 180° und 270° (-90°) Rotation.
    """
    try:
        data = request.get_json(force=True) or {}
        filename = data.get("filename")
        direction = data.get("direction", "right")  # "right" = 90°, "left" = -90°, "180" = 180°

        if not filename:
            return jsonify(success=False, message="Kein Dateiname angegeben"), 400

        # Datei im Staging oder INPUT_ROOT finden
        file_path = None
        in_staging = False

        # Zuerst im Staging suchen
        if fs.session_id:
            staged = os.path.join(fs.work_dir, filename)
            if os.path.exists(staged):
                file_path = staged
                in_staging = True

        # Dann im INPUT_ROOT suchen
        if not file_path:
            original = os.path.join(INPUT_ROOT, filename)
            if os.path.exists(original):
                # Kopiere Datei ins Staging vor der Rotation
                ensure_staging()
                import shutil
                staged_path = os.path.join(fs.work_dir, filename)
                shutil.copy2(original, staged_path)
                file_path = staged_path
                in_staging = True
                log(f"📋 Datei fürs Rotieren ins Staging kopiert: {filename}")

        if not file_path:
            log(f"❌ Datei nicht gefunden: {filename}", level="error")
            return jsonify(success=False, message="Datei nicht gefunden"), 404

        # Rotation bestimmen
        if direction == "180":
            angle = 180
        elif direction == "right":
            angle = 90
        else:  # left
            angle = -90

        ext = filename.lower().split('.')[-1] if '.' in filename else ''

        # PDF rotieren
        if ext == 'pdf':
            import fitz
            doc = fitz.open(file_path)
            for page in doc:
                page.set_rotation((page.rotation + angle) % 360)

            temp_path = file_path + ".tmp"
            doc.save(temp_path)
            doc.close()
            os.replace(temp_path, file_path)
            log(f"✅ PDF rotiert ({angle}°): {filename}")

        # Bilder rotieren (JPG, JPEG, PNG)
        elif ext in ['jpg', 'jpeg', 'png']:
            from PIL import Image
            img = Image.open(file_path)

            # PIL rotiert gegen den Uhrzeigersinn, wir wollen im Uhrzeigersinn
            # Also: 90° rechts = -90° in PIL
            pil_angle = -angle
            rotated = img.rotate(pil_angle, expand=True)

            # Format beibehalten
            if ext == 'png':
                rotated.save(file_path, 'PNG')
            else:
                # EXIF-Daten entfernen, da Rotation jetzt physisch ist
                rotated.save(file_path, 'JPEG', quality=95)

            img.close()
            rotated.close()
            log(f"✅ Bild rotiert ({angle}°): {filename}")

        else:
            return jsonify(success=False, message=f"Dateityp '{ext}' kann nicht rotiert werden"), 400

        return jsonify(
            success=True,
            message=f"Datei um {angle}° gedreht",
            filename=filename
        )

    except Exception as e:
        log(f"❌ rotate_file: {e}", level="error")
        return jsonify(success=False, message=str(e)), 500


@file_bp.route("/list_staged_files", methods=["GET"])
def list_staged_files():
    """
    Gibt alle Dateien im aktuellen Staging zurück.
    """
    try:
        if not fs.session_id:
            return jsonify(success=True, files=[])

        files = []
        staging_dir = fs.work_dir

        if staging_dir.exists():
            for item in staging_dir.rglob('*'):
                if item.is_file():
                    # Nur Dateiname ohne Pfad zurückgeben
                    files.append(item.name)

        log(f"📋 Staging-Dateien aufgelistet: {len(files)} Dateien")
        return jsonify(success=True, files=sorted(files))

    except Exception as e:
        log(f"❌ Fehler beim Auflisten der Staging-Dateien: {e}", level="error")
        return jsonify(success=False, message=str(e)), 500
