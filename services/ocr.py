import os
import re
import shutil
import subprocess
import img2pdf
import json
from .summarizer import summarize_pdf
from .file_utils import fs, to_rel_under_input, safe_line, handle_successful_processing, build_absender
from .logger import log
from config import INPUT_ROOT, FAIL_DIR_MEDIDOK, JSON_FOLDER
from flask import render_template, session

def process_medidok_files(file_paths, target_dir_unused):
    """
    Verarbeitet Medidok-Dateien staging-sicher:
    - Unterstützt Dateien aus INPUT_ROOT UND Staging
    - Bilddateien -> temp PDF im Staging -> OCR im Staging
    - KEIN Verschieben nach FAIL_DIR, KEIN Löschen der Originale
    - Rückgabe: Liste Ergebnis-Metadaten für UI
    """
    results = []

    for file_identifier in file_paths:
        # file_identifier kann sein:
        # - Nur Dateiname (z.B. "test.pdf")
        # - Relativer Pfad (z.B. "combined_20250114_123456.pdf")
        # - Absoluter Pfad (sollte nicht vorkommen)

        filename = os.path.basename(file_identifier)

        # ✅ WICHTIG: Original-Dateinamen JETZT speichern (ohne _ocr.pdf)
        true_original_filename = filename

        # Pfadauflösung: Staging ZUERST, dann Original
        working_input = None
        is_from_staging = False

        # 1. Versuch: Im Staging suchen
        if fs.session_id:
            staged_path = fs.work_dir / file_identifier
            if staged_path.exists():
                working_input = str(staged_path)
                is_from_staging = True
                log(f"📂 Staging-Datei gefunden: {filename}")

        # 2. Versuch: Im INPUT_ROOT suchen
        if not working_input:
            original_path = os.path.join(INPUT_ROOT, filename)
            if os.path.exists(original_path):
                working_input = original_path
                log(f"📂 Original-Datei gefunden: {filename}")

        # ✅ DOCX-Dateien: Zu PDF konvertieren und direkt analysieren (ohne OCR)
        if filename.lower().endswith('.docx'):
            log(f"📝 DOCX-Datei erkannt - konvertiere zu PDF (ohne OCR): {filename}")

            if not working_input:
                log(f"❌ DOCX-Datei nicht gefunden: {filename}", level="error")
                continue

            try:
                from docx import Document
                import fitz  # PyMuPDF

                # DOCX öffnen und Text extrahieren
                doc = Document(working_input)

                # Text aus allen Paragraphen sammeln
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)

                # Text aus Tabellen extrahieren
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text)
                        full_text.append(" | ".join(row_text))

                docx_content = "\n".join(full_text)

                # PDF erstellen mit PyMuPDF
                temp_rel = f"{os.path.splitext(filename)[0]}_docx_converted.pdf"
                temp_pdf = fs.work_dir / temp_rel
                temp_pdf.parent.mkdir(parents=True, exist_ok=True)

                pdf_doc = fitz.open()
                page = pdf_doc.new_page(width=595, height=842)  # A4-Format

                text_lines = docx_content.split('\n')
                y_position = 50
                line_height = 14
                max_y = 792

                for line in text_lines:
                    if y_position + line_height > max_y:
                        page = pdf_doc.new_page(width=595, height=842)
                        y_position = 50

                    try:
                        page.insert_text(
                            (50, y_position),
                            line,
                            fontsize=10,
                            fontname="helv"
                        )
                    except Exception:
                        try:
                            ascii_line = line.encode('ascii', 'replace').decode('ascii')
                            page.insert_text(
                                (50, y_position),
                                ascii_line,
                                fontsize=10,
                                fontname="helv"
                            )
                        except:
                            pass

                    y_position += line_height

                page_count = len(pdf_doc)
                pdf_doc.save(str(temp_pdf))
                pdf_doc.close()

                log(f"✅ DOCX zu PDF konvertiert ({len(text_lines)} Zeilen, {page_count} Seiten): {temp_rel}")

                # PDF analysieren (ohne OCR)
                summary = summarize_pdf(str(temp_pdf))
                text = (summary or "").strip()
                lines = [(s or "").strip() for s in text.split("\n")]

                # Neue 8-Zeilen-Logik
                log(f"🔍 [DEBUG] LLM-Ausgabe hat {len(lines)} Zeilen")
                for i, line in enumerate(lines[:8]):
                    log(f"   Zeile {i+1}: '{line}'")

                fachrichtung = safe_line(lines, 4, "")
                name_absender = safe_line(lines, 5, "")
                category = safe_line(lines, 7, "11")
                absender = build_absender(fachrichtung, name_absender, category)

                log(f"📋 [DEBUG] Extrahiert:")
                log(f"   Fachrichtung: '{fachrichtung}'")
                log(f"   Name: '{name_absender}'")
                log(f"   Kategorie: '{category}'")
                log(f"   → Absender: '{absender}'")

                summary_data = {
                    "file": filename,
                    "filename": filename,
                    "originalFilename": true_original_filename,
                    "name": safe_line(lines, 0, "Unbekannt"),
                    "vorname": safe_line(lines, 1, "Unbekannt"),
                    "geburtsdatum": safe_line(lines, 2, "Unbekannt"),
                    "datum": safe_line(lines, 3, "Unbekannt"),
                    "beschreibung1": absender,
                    "beschreibung2": safe_line(lines, 6, "Keine Beschreibung verfügbar"),
                    "categoryID": category,
                }

                result = {"summary": summary_data}
                log(f"✅ DOCX-Datei analysiert: {filename}")
                log(f"📋 Result: {result}")
                results.append(result)
                continue

            except Exception as e:
                log(f"❌ Fehler bei DOCX-Analyse: {e}", level="error")
                continue

        # ✅ TXT-Dateien: Direkt analysieren ohne OCR
        if filename.lower().endswith('.txt'):
            log(f"📝 TXT-Datei erkannt - überspringe OCR: {filename}")

            if not working_input:
                log(f"❌ TXT-Datei nicht gefunden: {filename}", level="error")
                continue

            # TXT-Inhalt direkt lesen
            try:
                encodings = ['utf-8', 'windows-1252', 'latin-1', 'iso-8859-1', 'cp1252']
                txt_content = None

                for encoding in encodings:
                    try:
                        with open(working_input, 'r', encoding=encoding) as f:
                            txt_content = f.read()
                        log(f"✅ TXT-Datei gelesen mit {encoding}: {filename}")
                        break
                    except (UnicodeDecodeError, LookupError):
                        continue

                if txt_content is None:
                    with open(working_input, 'rb') as f:
                        raw_content = f.read()
                    txt_content = raw_content.decode('utf-8', errors='replace')
                    log(f"⚠️ TXT-Datei mit errors='replace' gelesen: {filename}")

                # TXT zu temporärem PDF konvertieren für einheitliche Verarbeitung
                temp_rel = f"{os.path.splitext(filename)[0]}_txt_converted.pdf"
                temp_pdf = fs.work_dir / temp_rel
                temp_pdf.parent.mkdir(parents=True, exist_ok=True)

                # Text als einfaches PDF erstellen mit PyMuPDF (fitz)
                import fitz  # PyMuPDF ist bereits installiert

                doc = fitz.open()  # Neues leeres PDF
                page = doc.new_page(width=595, height=842)  # A4-Format

                # Text zeilenweise einfügen (robuster als textbox)
                text_lines = txt_content.split('\n')
                y_position = 50
                line_height = 14
                max_y = 792  # Unterer Rand
                page_num = 1

                for line in text_lines:
                    # Neue Seite falls nötig
                    if y_position + line_height > max_y:
                        page = doc.new_page(width=595, height=842)
                        y_position = 50
                        page_num += 1

                    # Zeile einfügen
                    try:
                        page.insert_text(
                            (50, y_position),
                            line,
                            fontsize=10,
                            fontname="helv"
                        )
                    except Exception as e:
                        # Bei Fehler: Zeile mit ASCII-only versuchen
                        try:
                            ascii_line = line.encode('ascii', 'replace').decode('ascii')
                            page.insert_text(
                                (50, y_position),
                                ascii_line,
                                fontsize=10,
                                fontname="helv"
                            )
                        except:
                            log(f"⚠️ Zeile konnte nicht eingefügt werden: {str(e)[:50]}")

                    y_position += line_height

                page_count = len(doc)  # Vor dem Schließen speichern
                doc.save(str(temp_pdf))
                doc.close()

                log(f"✅ TXT zu PDF konvertiert ({len(text_lines)} Zeilen, {page_count} Seiten): {temp_rel}")

                # Jetzt normale PDF-Analyse mit bewährter Funktion
                summary = summarize_pdf(str(temp_pdf))
                text = (summary or "").strip()
                lines = [(s or "").strip() for s in text.split("\n")]

                # Neue 8-Zeilen-Logik
                log(f"🔍 [DEBUG] LLM-Ausgabe hat {len(lines)} Zeilen")
                for i, line in enumerate(lines[:8]):
                    log(f"   Zeile {i+1}: '{line}'")

                fachrichtung = safe_line(lines, 4, "")
                name_absender = safe_line(lines, 5, "")
                category = safe_line(lines, 7, "11")
                absender = build_absender(fachrichtung, name_absender, category)

                log(f"📋 [DEBUG] Extrahiert:")
                log(f"   Fachrichtung: '{fachrichtung}'")
                log(f"   Name: '{name_absender}'")
                log(f"   Kategorie: '{category}'")
                log(f"   → Absender: '{absender}'")

                summary_data = {
                    "file": filename,
                    "filename": filename,
                    "originalFilename": true_original_filename,
                    "name": safe_line(lines, 0, "Unbekannt"),
                    "vorname": safe_line(lines, 1, "Unbekannt"),
                    "geburtsdatum": safe_line(lines, 2, "Unbekannt"),
                    "datum": safe_line(lines, 3, "Unbekannt"),
                    "beschreibung1": absender,
                    "beschreibung2": safe_line(lines, 6, "Keine Beschreibung verfügbar"),
                    "categoryID": category,
                }

                result = {"summary": summary_data}
                log(f"✅ TXT-Datei analysiert: {filename}")
                log(f"📋 Result: {result}")
                results.append(result)
                continue

            except Exception as e:
                log(f"❌ Fehler bei TXT-Analyse: {e}", level="error")
                continue

        # ✅ Dateien ohne gültige Dateiendung als .jpg behandeln und umbenennen
        has_valid_extension = filename.lower().endswith(('.pdf', '.jpg', '.jpeg', '.png', '.tif', '.tiff', '.docx'))
        if working_input and not has_valid_extension:
            log(f"📝 Datei ohne gültige Endung erkannt: {filename}")
            new_filename = filename + '.jpg'

            if is_from_staging:
                # Im Staging umbenennen
                new_path = fs.work_dir / new_filename
                import shutil
                shutil.move(working_input, str(new_path))
                working_input = str(new_path)
                filename = new_filename
                true_original_filename = new_filename
                log(f"✅ Im Staging umbenannt: {filename}")
            else:
                # Im INPUT_ROOT umbenennen
                new_path = os.path.join(INPUT_ROOT, new_filename)
                import shutil
                shutil.move(working_input, new_path)
                working_input = new_path
                filename = new_filename
                true_original_filename = new_filename
                log(f"✅ In INPUT_ROOT umbenannt: {filename}")

        if not filename.lower().endswith(('.pdf', '.jpg', '.jpeg', '.png', '.docx')):
            log(f"⏭️ Überspringe Datei (kein unterstütztes Format): {filename}")
            continue
        
        # 3. Nicht gefunden
        if not working_input:
            log(f"❌ Datei nicht gefunden (weder Staging noch Input): {filename}", level="error")
            continue

        # Falls Bild: erst staging-Temp-PDF erzeugen
        if filename.lower().endswith(('.jpg', '.jpeg', '.png')):
            temp_rel = f"{os.path.splitext(filename)[0]}_converted.pdf"
            temp_pdf = fs.work_dir / temp_rel
            temp_pdf.parent.mkdir(parents=True, exist_ok=True)
            
            with open(temp_pdf, "wb") as f:
                f.write(img2pdf.convert([working_input]))
            
            working_input = str(temp_pdf)
            log(f"🖼️ Bild zu PDF konvertiert: {temp_rel}")

        # OCR im Staging
        base_no_ext = os.path.splitext(filename)[0]
        
        # ✅ WICHTIG: out_rel sollte FLACH sein (kein Unterverzeichnis)
        out_rel = f"{base_no_ext}_ocr.pdf"
        
        staged_out = ocr_to_staging(working_input, out_rel)

        if not staged_out:
            log(f"❌ Fehler bei Verarbeitung – Datei bleibt unverändert: {filename}")
            continue

        log(f"✅ OCR erstellt im Staging: {staged_out}")

        # Summary + Umbenennungsplan (handle_successful_processing plant nur!)
        summary = summarize_pdf(staged_out)
        text = (summary or "").strip()
        lines = [(s or "").strip() for s in text.split("\n")]

        # Neue 8-Zeilen-Logik
        log(f"🔍 [DEBUG] LLM-Ausgabe hat {len(lines)} Zeilen")
        for i, line in enumerate(lines[:8]):
            log(f"   Zeile {i+1}: '{line}'")

        fachrichtung = safe_line(lines, 4, "")
        name_absender = safe_line(lines, 5, "")
        category = safe_line(lines, 7, "11")
        absender = build_absender(fachrichtung, name_absender, category)

        log(f"📋 [DEBUG] Extrahiert:")
        log(f"   Fachrichtung: '{fachrichtung}'")
        log(f"   Name: '{name_absender}'")
        log(f"   Kategorie: '{category}'")
        log(f"   → Absender: '{absender}'")

        summary_data = {
            "file": out_rel,  # ✅ Relativer Pfad für Frontend
            "filename": os.path.basename(staged_out),
            "originalFilename": true_original_filename,  # ✅ HIER: Echtes Original verwenden!
            "name": safe_line(lines, 0, "Unbekannt"),
            "vorname": safe_line(lines, 1, "Unbekannt"),
            "geburtsdatum": safe_line(lines, 2, "Unbekannt"),
            "datum": safe_line(lines, 3, "Unbekannt"),
            "beschreibung1": absender,
            "beschreibung2": safe_line(lines, 6, "Keine Beschreibung verfügbar"),
            "categoryID": category,
        }
        
        create_control_json_from_summaries([summary_data])

        result = handle_successful_processing(
            summary_data=summary_data,
            original_path=staged_out,  # hier: die STAGING-Datei (absoluter Pfad)
            target_dir=os.path.dirname(staged_out)
        )

        # Aktualisiere summary_data mit dem neuen Dateinamen
        summary_data["file"] = result["renamed"]

        # Control JSON mit aktualisiertem Dateinamen neu schreiben (mit originalFilename als Key für Dedupe)
        create_control_json_from_summaries([summary_data], key="originalFilename")

        log(f"📋 Result: {result}")
        results.append(result)

    return results

def process_medidok_files_with_model(file_paths, target_dir_unused, model, session_id):
    """
    Thread-sichere Variante von process_medidok_files für Background-Threads.
    Benötigt explizites Modell und Session-ID (kein Flask Request-Context).

    Args:
        file_paths: Liste von Dateipfaden
        target_dir_unused: Wird nicht verwendet (Kompatibilität)
        model: LLM-Modell explizit (z.B. "mistral-nemo:latest")
        session_id: Flask Session-ID explizit

    Returns:
        list: Liste von Result-Dictionaries mit 'summary' Key
    """
    from .summarizer import summarize_pdf

    results = []

    for file_identifier in file_paths:
        filename = os.path.basename(file_identifier)

        true_original_filename = filename

        # Pfadauflösung: Staging ZUERST, dann Original
        working_input = None
        is_from_staging = False

        if fs.session_id:
            staged_path = fs.work_dir / file_identifier
            if staged_path.exists():
                working_input = str(staged_path)
                is_from_staging = True
                log(f"📂 Staging-Datei gefunden: {filename}")

        if not working_input:
            original_path = os.path.join(INPUT_ROOT, filename)
            if os.path.exists(original_path):
                working_input = original_path
                log(f"📂 Original-Datei gefunden: {filename}")

        # ✅ DOCX-Dateien: Zu PDF konvertieren und direkt analysieren (ohne OCR)
        if filename.lower().endswith('.docx'):
            log(f"📝 DOCX-Datei erkannt - konvertiere zu PDF (ohne OCR, Modell: {model}): {filename}")

            if not working_input:
                log(f"❌ DOCX-Datei nicht gefunden: {filename}", level="error")
                continue

            try:
                from docx import Document
                import fitz  # PyMuPDF

                # DOCX öffnen und Text extrahieren
                doc = Document(working_input)

                # Text aus allen Paragraphen sammeln
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)

                # Text aus Tabellen extrahieren
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text)
                        full_text.append(" | ".join(row_text))

                docx_content = "\n".join(full_text)

                # PDF erstellen mit PyMuPDF
                temp_rel = f"{os.path.splitext(filename)[0]}_docx_converted.pdf"
                temp_pdf = fs.work_dir / temp_rel
                temp_pdf.parent.mkdir(parents=True, exist_ok=True)

                pdf_doc = fitz.open()
                page = pdf_doc.new_page(width=595, height=842)  # A4-Format

                text_lines = docx_content.split('\n')
                y_position = 50
                line_height = 14
                max_y = 792

                for line in text_lines:
                    if y_position + line_height > max_y:
                        page = pdf_doc.new_page(width=595, height=842)
                        y_position = 50

                    try:
                        page.insert_text(
                            (50, y_position),
                            line,
                            fontsize=10,
                            fontname="helv"
                        )
                    except Exception:
                        try:
                            ascii_line = line.encode('ascii', 'replace').decode('ascii')
                            page.insert_text(
                                (50, y_position),
                                ascii_line,
                                fontsize=10,
                                fontname="helv"
                            )
                        except:
                            pass

                    y_position += line_height

                page_count = len(pdf_doc)
                pdf_doc.save(str(temp_pdf))
                pdf_doc.close()

                log(f"✅ DOCX zu PDF konvertiert ({len(text_lines)} Zeilen, {page_count} Seiten): {temp_rel}")

                # PDF analysieren (ohne OCR, mit explizitem Modell)
                summary = summarize_pdf(str(temp_pdf), model=model)
                text = (summary or "").strip()
                lines = [(s or "").strip() for s in text.split("\n")]

                # Neue 8-Zeilen-Logik
                log(f"🔍 [DEBUG] LLM-Ausgabe hat {len(lines)} Zeilen")
                for i, line in enumerate(lines[:8]):
                    log(f"   Zeile {i+1}: '{line}'")

                fachrichtung = safe_line(lines, 4, "")
                name_absender = safe_line(lines, 5, "")
                category = safe_line(lines, 7, "11")
                absender = build_absender(fachrichtung, name_absender, category)

                log(f"📋 [DEBUG] Extrahiert:")
                log(f"   Fachrichtung: '{fachrichtung}'")
                log(f"   Name: '{name_absender}'")
                log(f"   Kategorie: '{category}'")
                log(f"   → Absender: '{absender}'")

                summary_data = {
                    "file": filename,
                    "filename": filename,
                    "originalFilename": true_original_filename,
                    "name": safe_line(lines, 0, "Unbekannt"),
                    "vorname": safe_line(lines, 1, "Unbekannt"),
                    "geburtsdatum": safe_line(lines, 2, "Unbekannt"),
                    "datum": safe_line(lines, 3, "Unbekannt"),
                    "beschreibung1": absender,
                    "beschreibung2": safe_line(lines, 6, "Keine Beschreibung verfügbar"),
                    "categoryID": category,
                }

                result = {"summary": summary_data}
                log(f"✅ DOCX-Datei analysiert (mit Modell {model}): {filename}")
                log(f"📋 Result: {result}")
                results.append(result)
                continue

            except Exception as e:
                log(f"❌ Fehler bei DOCX-Analyse: {e}", level="error")
                import traceback
                log(traceback.format_exc(), level="error")
                continue

        # ✅ TXT-Dateien: Direkt analysieren ohne OCR
        if filename.lower().endswith('.txt'):
            log(f"📝 TXT-Datei erkannt - überspringe OCR (mit explizitem Modell): {filename}")

            if not working_input:
                log(f"❌ TXT-Datei nicht gefunden: {filename}", level="error")
                continue

            # TXT-Inhalt direkt lesen
            try:
                encodings = ['utf-8', 'windows-1252', 'latin-1', 'iso-8859-1', 'cp1252']
                txt_content = None

                for encoding in encodings:
                    try:
                        with open(working_input, 'r', encoding=encoding) as f:
                            txt_content = f.read()
                        log(f"✅ TXT-Datei gelesen mit {encoding}: {filename}")
                        break
                    except (UnicodeDecodeError, LookupError):
                        continue

                if txt_content is None:
                    with open(working_input, 'rb') as f:
                        raw_content = f.read()
                    txt_content = raw_content.decode('utf-8', errors='replace')
                    log(f"⚠️ TXT-Datei mit errors='replace' gelesen: {filename}")

                # TXT zu temporärem PDF konvertieren für einheitliche Verarbeitung
                temp_rel = f"{os.path.splitext(filename)[0]}_txt_converted.pdf"
                temp_pdf = fs.work_dir / temp_rel
                temp_pdf.parent.mkdir(parents=True, exist_ok=True)

                # Text als einfaches PDF erstellen mit PyMuPDF (fitz)
                import fitz  # PyMuPDF ist bereits installiert

                doc = fitz.open()  # Neues leeres PDF
                page = doc.new_page(width=595, height=842)  # A4-Format

                # Text zeilenweise einfügen (robuster als textbox)
                text_lines = txt_content.split('\n')
                y_position = 50
                line_height = 14
                max_y = 792  # Unterer Rand
                page_num = 1

                for line in text_lines:
                    # Neue Seite falls nötig
                    if y_position + line_height > max_y:
                        page = doc.new_page(width=595, height=842)
                        y_position = 50
                        page_num += 1

                    # Zeile einfügen
                    try:
                        page.insert_text(
                            (50, y_position),
                            line,
                            fontsize=10,
                            fontname="helv"
                        )
                    except Exception as e:
                        # Bei Fehler: Zeile mit ASCII-only versuchen
                        try:
                            ascii_line = line.encode('ascii', 'replace').decode('ascii')
                            page.insert_text(
                                (50, y_position),
                                ascii_line,
                                fontsize=10,
                                fontname="helv"
                            )
                        except:
                            log(f"⚠️ Zeile konnte nicht eingefügt werden: {str(e)[:50]}")

                    y_position += line_height

                page_count = len(doc)  # Vor dem Schließen speichern
                doc.save(str(temp_pdf))
                doc.close()

                log(f"✅ TXT zu PDF konvertiert ({len(text_lines)} Zeilen, {page_count} Seiten, Modell: {model}): {temp_rel}")

                # Jetzt normale PDF-Analyse mit bewährter Funktion
                summary = summarize_pdf(str(temp_pdf), model=model)
                text = (summary or "").strip()
                lines = [(s or "").strip() for s in text.split("\n")]

                # Neue 8-Zeilen-Logik
                log(f"🔍 [DEBUG] LLM-Ausgabe hat {len(lines)} Zeilen")
                for i, line in enumerate(lines[:8]):
                    log(f"   Zeile {i+1}: '{line}'")

                fachrichtung = safe_line(lines, 4, "")
                name_absender = safe_line(lines, 5, "")
                category = safe_line(lines, 7, "11")
                absender = build_absender(fachrichtung, name_absender, category)

                log(f"📋 [DEBUG] Extrahiert:")
                log(f"   Fachrichtung: '{fachrichtung}'")
                log(f"   Name: '{name_absender}'")
                log(f"   Kategorie: '{category}'")
                log(f"   → Absender: '{absender}'")

                summary_data = {
                    "file": filename,
                    "filename": filename,
                    "originalFilename": true_original_filename,
                    "name": safe_line(lines, 0, "Unbekannt"),
                    "vorname": safe_line(lines, 1, "Unbekannt"),
                    "geburtsdatum": safe_line(lines, 2, "Unbekannt"),
                    "datum": safe_line(lines, 3, "Unbekannt"),
                    "beschreibung1": absender,
                    "beschreibung2": safe_line(lines, 6, "Keine Beschreibung verfügbar"),
                    "categoryID": category,
                }

                result = {"summary": summary_data}
                log(f"✅ TXT-Datei analysiert (mit Modell {model}): {filename}")
                log(f"📋 Result: {result}")
                results.append(result)
                continue

            except Exception as e:
                log(f"❌ Fehler bei TXT-Analyse: {e}", level="error")
                import traceback
                log(traceback.format_exc(), level="error")
                continue

        # ✅ Dateien ohne gültige Dateiendung als .jpg behandeln und umbenennen
        has_valid_extension = filename.lower().endswith(('.pdf', '.jpg', '.jpeg', '.png', '.tif', '.tiff', '.docx'))
        if working_input and not has_valid_extension:
            log(f"📝 Datei ohne gültige Endung erkannt: {filename}")
            new_filename = filename + '.jpg'

            if is_from_staging:
                # Im Staging umbenennen
                new_path = fs.work_dir / new_filename
                shutil.move(working_input, str(new_path))
                working_input = str(new_path)
                filename = new_filename
                true_original_filename = new_filename
                log(f"✅ Im Staging umbenannt: {filename}")
            else:
                # Im INPUT_ROOT umbenennen
                new_path = os.path.join(INPUT_ROOT, new_filename)
                shutil.move(working_input, new_path)
                working_input = new_path
                filename = new_filename
                true_original_filename = new_filename
                log(f"✅ In INPUT_ROOT umbenannt: {filename}")

        if not filename.lower().endswith(('.pdf', '.jpg', '.jpeg', '.png', '.docx')):
            log(f"⏭️ Überspringe Datei (kein unterstütztes Format): {filename}")
            continue

        if not working_input:
            log(f"❌ Datei nicht gefunden: {filename}", level="error")
            continue

        # Falls Bild: erst staging-Temp-PDF erzeugen
        if filename.lower().endswith(('.jpg', '.jpeg', '.png')):
            temp_rel = f"{os.path.splitext(filename)[0]}_converted.pdf"
            temp_pdf = fs.work_dir / temp_rel
            temp_pdf.parent.mkdir(parents=True, exist_ok=True)
            
            with open(temp_pdf, "wb") as f:
                f.write(img2pdf.convert([working_input]))
            
            working_input = str(temp_pdf)
            log(f"🖼️ Bild zu PDF konvertiert: {temp_rel}")

        # OCR im Staging
        base_no_ext = os.path.splitext(filename)[0]
        out_rel = f"{base_no_ext}_ocr.pdf"
        
        staged_out = ocr_to_staging(working_input, out_rel)

        if not staged_out:
            log(f"❌ Fehler bei Verarbeitung: {filename}")
            continue

        log(f"✅ OCR erstellt im Staging: {staged_out}")

        # ✅ Summary mit explizitem Modell (thread-sicher)
        summary = summarize_pdf(staged_out, model=model)
        text = (summary or "").strip()
        lines = [(s or "").strip() for s in text.split("\n")]

        # Neue 8-Zeilen-Logik
        log(f"🔍 [DEBUG] LLM-Ausgabe hat {len(lines)} Zeilen")
        for i, line in enumerate(lines[:8]):
            log(f"   Zeile {i+1}: '{line}'")

        fachrichtung = safe_line(lines, 4, "")
        name_absender = safe_line(lines, 5, "")
        category = safe_line(lines, 7, "11")
        absender = build_absender(fachrichtung, name_absender, category)

        log(f"📋 [DEBUG] Extrahiert:")
        log(f"   Fachrichtung: '{fachrichtung}'")
        log(f"   Name: '{name_absender}'")
        log(f"   Kategorie: '{category}'")
        log(f"   → Absender: '{absender}'")

        summary_data = {
            "file": out_rel,
            "filename": os.path.basename(staged_out),
            "originalFilename": true_original_filename,
            "name": safe_line(lines, 0, "Unbekannt"),
            "vorname": safe_line(lines, 1, "Unbekannt"),
            "geburtsdatum": safe_line(lines, 2, "Unbekannt"),
            "datum": safe_line(lines, 3, "Unbekannt"),
            "beschreibung1": absender,
            "beschreibung2": safe_line(lines, 6, "Keine Beschreibung verfügbar"),
            "categoryID": category,
        }

        result = handle_successful_processing(
            summary_data=summary_data,
            original_path=staged_out,
            target_dir=os.path.dirname(staged_out)
        )

        # Aktualisiere summary_data mit dem neuen Dateinamen
        summary_data["file"] = result["renamed"]

        log(f"📋 Result: {result}")
        results.append(result)

    return results

def create_control_json_from_summaries(summaries, *, overwrite=False, dedupe=True, key="file"):
    """Schreibt/aktualisiert control_<session>.json.
    - overwrite=True: bestehende Datei wird ersetzt
    - dedupe=True: doppelte Einträge (gleicher key) werden überschrieben
    """
    sid = session.get("session_id", "default")
    path = os.path.join(JSON_FOLDER, f"control_{sid}.json")
    os.makedirs(JSON_FOLDER, exist_ok=True)

    # Basisdaten laden
    if not overwrite and os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                control_data = json.load(f)
            if not isinstance(control_data, list):
                control_data = []
        except Exception:
            control_data = []
    else:
        control_data = []

    # Dedupe-Map aufbauen - NUTZE ORIGINALFILENAME als eindeutigen Schlüssel
    if dedupe and control_data:
        # Verwende originalFilename statt des konfigurierbaren key-Parameters für Eindeutigkeit
        index = {str(entry.get("originalFilename", "")): i for i, entry in enumerate(control_data)}
    else:
        index = {}

    # Neue Einträge einpflegen
    for s in summaries:
        entry = {
            "file": s.get("file",""),
            "filename": s.get("filename",""),
            "originalFilename": s.get("originalFilename", s.get("filename","")),  # ✅ WICHTIG: Original-Namen bewahren!
            "name": s.get("name",""),
            "vorname": s.get("vorname",""),
            "geburtsdatum": s.get("geburtsdatum",""),
            "datum": s.get("datum",""),
            "beschreibung1": s.get("beschreibung1",""),
            "beschreibung2": s.get("beschreibung2",""),
            "categoryID": s.get("categoryID",""),
            "selected": True,
        }
        # Nutze originalFilename als eindeutigen Schlüssel (nicht den key-Parameter)
        k = str(entry.get("originalFilename", ""))
        if dedupe and k and k in index:
            control_data[index[k]] = entry  # überschreiben (neueste Werte)
        else:
            index[k] = len(control_data)
            control_data.append(entry)

    with open(path, "w", encoding="utf-8") as f:
        json.dump(control_data, f, ensure_ascii=False, indent=2)

    log(f"[INFO] control.json aktualisiert ({len(control_data)} Einträge): {path}")

def ocr_to_staging(input_pdf_path: str, output_rel: str):
    output_basename = os.path.basename(output_rel)
    staged_out = os.path.join(fs.work_dir, output_basename)

    os.makedirs(os.path.dirname(staged_out), exist_ok=True)

    log(f"🔍 [DEBUG] OCR-Start: {os.path.basename(input_pdf_path)} → {output_basename}")

    try:
        result = subprocess.run(
            ['ocrmypdf', '-l', 'deu', '--force-ocr', '--deskew', '--rotate-pages', '--clean', '-O', '0', '--invalidate-digital-signatures', input_pdf_path, staged_out],
            check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=300
        )
        log(f"✅ [DEBUG] OCR erfolgreich: {output_basename}")

        # Log STDOUT und STDERR für Debugging
        if result.stdout:
            log(f"📝 [DEBUG] OCR stdout: {result.stdout.decode()[:200]}")
        if result.stderr:
            log(f"📝 [DEBUG] OCR stderr: {result.stderr.decode()[:200]}")

        # Prüfe Dateigröße
        if os.path.exists(staged_out):
            file_size = os.path.getsize(staged_out)
            log(f"📊 [DEBUG] OCR-Ausgabe: {file_size} Bytes")
    except subprocess.CalledProcessError as e:
        log(f"❌ OCR-Fehler bei {input_pdf_path}: {e.stderr.decode()}", level="error")
        return None

    if not os.path.exists(staged_out):
        log(f"❌ OCR-Zieldatei fehlt: {staged_out}", level="error")
        return None

    return staged_out